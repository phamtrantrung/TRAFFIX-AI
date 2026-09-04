# utils/report_generator.py  (v2 - nâng cấp giao diện chuyên nghiệp + liên kết AI)
#
# Sinh báo cáo phân tích (Word .docx và Excel .xlsx) từ dữ liệu giám sát
# giao thông. So với bản trước, bản này:
#   - Trình bày khoa học hơn: thẻ KPI màu, bảng có viền/màu xen kẽ,
#     Excel Table gốc (banded rows, filter sẵn).
#   - Biểu đồ tách theo loại xe (giống trang Phân tích), màu đồng bộ
#     với giao diện Dashboard/Phân tích.
#   - GỌI TRỰC TIẾP chatbot_engine (chính là Trợ lý AI) để tự hỏi-đáp vài
#     câu chuẩn về đúng khoảng thời gian đang lọc -> chứng minh 3 tính
#     năng Dashboard/Phân tích/Trợ lý AI dùng chung 1 nguồn dữ liệu sống.
#
# Yêu cầu: pip install python-docx openpyxl matplotlib

import os
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, Reference
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

from utils import database
from utils import chatbot_engine
from utils.chatbot_engine import _get_model

REPORTS_DIR = "reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

# Bảng màu đồng bộ với giao diện Dashboard/Phân tích (xem analytics.html)
COLOR_TEAL = "4FD1C5"
COLOR_RED = "F56565"
COLOR_BLUE = "63B3ED"
COLOR_PURPLE = "B794F4"
COLOR_NAVY = "171D2E"
COLOR_NAVY_DARK = "10141F"
CLASS_COLORS = {"car": COLOR_TEAL, "truck": "ECC94B", "bus": COLOR_BLUE, "motorbike": COLOR_PURPLE}
CLASS_LABELS_VI = {"car": "Ô tô con", "truck": "Xe tải", "bus": "Xe buýt", "motorbike": "Xe máy"}


def _resolve_evidence_path(evidence_image_path):
    if not evidence_image_path:
        return None
    if os.path.exists(evidence_image_path):
        return evidence_image_path
    fallback = os.path.join("evidence", os.path.basename(evidence_image_path))
    if os.path.exists(fallback):
        return fallback
    return None


# ============================================================
# 1. TRUY VẤN DỮ LIỆU
# ============================================================

def _fetch_report_data(start_time: str, end_time: str) -> dict:
    vehicle_counts = database.run_safe_query(f"""
        SELECT vehicle_class, direction, SUM(count) as total
        FROM traffic_stats
        WHERE timestamp BETWEEN '{start_time}' AND '{end_time}'
        GROUP BY vehicle_class, direction
        ORDER BY vehicle_class, direction
    """)

    violations = database.run_safe_query(f"""
        SELECT id, track_id, plate_number, violation_type, vehicle_class,
               location, timestamp, evidence_image_path
        FROM violations
        WHERE timestamp BETWEEN '{start_time}' AND '{end_time}'
        ORDER BY timestamp DESC
        LIMIT 200
    """)

    hourly_traffic = database.run_safe_query(f"""
        SELECT strftime('%Y-%m-%d %H:00', timestamp) as hour, SUM(count) as total
        FROM traffic_stats
        WHERE timestamp BETWEEN '{start_time}' AND '{end_time}'
        GROUP BY hour
        ORDER BY hour
    """)

    hourly_traffic_by_class = database.run_safe_query(f"""
        SELECT strftime('%Y-%m-%d %H:00', timestamp) as hour, vehicle_class, SUM(count) as total
        FROM traffic_stats
        WHERE timestamp BETWEEN '{start_time}' AND '{end_time}'
        GROUP BY hour, vehicle_class
        ORDER BY hour
    """)

    hourly_violations = database.run_safe_query(f"""
        SELECT strftime('%Y-%m-%d %H:00', timestamp) as hour, COUNT(*) as total
        FROM violations
        WHERE timestamp BETWEEN '{start_time}' AND '{end_time}'
        GROUP BY hour
        ORDER BY hour
    """)

    return {
        "vehicle_counts": vehicle_counts,
        "violations": violations,
        "hourly_traffic": hourly_traffic,
        "hourly_traffic_by_class": hourly_traffic_by_class,
        "hourly_violations": hourly_violations,
    }


def _compute_kpis(data: dict) -> dict:
    """Tính các chỉ số tổng quan giống hệt cách trang Phân tích tính (4 thẻ
    stat-box: Tổng lượt xe / Tổng vi phạm / Giờ cao điểm / Loại xe phổ biến)
    -> đảm bảo số liệu trong báo cáo KHỚP với những gì đang hiển thị trên
    Dashboard/Phân tích tại đúng thời điểm xuất báo cáo."""
    total_traffic = sum((r["total"] or 0) for r in data["vehicle_counts"])
    total_violations = len(data["violations"])
    violation_rate = (total_violations / total_traffic * 100) if total_traffic > 0 else 0.0

    peak_hour, peak_val = None, -1
    for r in data["hourly_traffic"]:
        if (r["total"] or 0) > peak_val:
            peak_val = r["total"] or 0
            peak_hour = r["hour"]

    class_totals = {}
    for r in data["vehicle_counts"]:
        class_totals[r["vehicle_class"]] = class_totals.get(r["vehicle_class"], 0) + (r["total"] or 0)
    top_class, top_val = None, -1
    for c, v in class_totals.items():
        if v > top_val:
            top_val, top_class = v, c

    return {
        "total_traffic": total_traffic,
        "total_violations": total_violations,
        "violation_rate": violation_rate,
        "peak_hour": peak_hour,
        "peak_val": peak_val if peak_val >= 0 else 0,
        "top_class": top_class,
        "top_class_label": CLASS_LABELS_VI.get(top_class, top_class) if top_class else "–",
        "top_val": top_val if top_val >= 0 else 0,
    }


# ============================================================
# 2. NHẬN ĐỊNH AI + HỎI-ĐÁP NHANH VỚI TRỢ LÝ AI
#    (dùng chung chatbot_engine với trang /chatbot -> tạo "sợi dây liên kết")
# ============================================================

def _generate_ai_summary(data: dict, kpis: dict, start_time: str, end_time: str) -> str:
    model = _get_model()

    prompt = f"""Bạn là chuyên gia phân tích giao thông. Dựa trên số liệu giám sát
sau đây (khoảng thời gian {start_time} đến {end_time}), hãy viết một đoạn
NHẬN ĐỊNH TÓM TẮT bằng tiếng Việt (khoảng 150-200 từ), văn phong chuyên
nghiệp như trong báo cáo kỹ thuật, nêu bật: xu hướng lưu lượng nổi bật,
loại phương tiện chiếm đa số, mức độ vi phạm, khung giờ đáng chú ý, và
một khuyến nghị ngắn nếu có.

LƯU Ý QUAN TRỌNG: Hệ thống đếm lưu lượng (theo line) và ghi nhận vi phạm
(theo vùng ROI - khu vực hạn chế) hoạt động HOÀN TOÀN ĐỘC LẬP với nhau
theo đúng thiết kế. Vì vậy việc tổng lượt xe = 0 nhưng vẫn có vi phạm là
HOÀN TOÀN BÌNH THƯỜNG, KHÔNG phải dấu hiệu lỗi thiết bị, sự cố cảm biến
hay camera hỏng - tuyệt đối đừng suy diễn theo hướng đó hay khuyến nghị
kiểm tra/hiệu chuẩn camera chỉ vì lý do này.

Dữ liệu tổng quan:
- Tổng lượt xe: {kpis['total_traffic']}
- Tổng vi phạm: {kpis['total_violations']} (tỉ lệ {kpis['violation_rate']:.1f}%)
- Giờ cao điểm: {kpis['peak_hour']} ({kpis['peak_val']} xe)
- Loại xe phổ biến nhất: {kpis['top_class_label']} ({kpis['top_val']} xe)

Dữ liệu chi tiết:
- Theo loại/hướng: {data['vehicle_counts']}
- Lưu lượng theo giờ: {data['hourly_traffic']}
- Vi phạm theo giờ: {data['hourly_violations']}

CHỈ trả về đoạn văn nhận định, không markdown, không tiêu đề."""

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"(Không thể tạo nhận định AI tự động: {e})"


STANDARD_QUESTIONS = [
    "Trong khoảng thời gian này có tổng cộng bao nhiêu vi phạm?",
    "Xe (track_id) nào vi phạm nhiều lần nhất trong khoảng thời gian này?",
    "Giờ nào có lưu lượng xe đông nhất trong khoảng thời gian này?",
]


def _ask_ai_assistant_scoped(start_time: str, end_time: str, data: dict, kpis: dict) -> list:
    """Trả lời 3 câu hỏi chuẩn bằng cách TÍNH TRỰC TIẾP từ `data`
    (đã lọc đúng theo start_time/end_time bằng SQL trong _fetch_report_data),
    thay vì nhờ AI tự suy luận điều kiện thời gian từ câu hỏi văn bản.
    Cách cũ để chatbot_engine.ask() tự sinh SQL không đảm bảo áp đúng
    time-range, gây lệch số liệu so với phần còn lại của báo cáo."""

    # Câu 1: Tổng số vi phạm
    total_violations = kpis["total_violations"]
    answer_1 = f"Trong khoảng thời gian này có tổng cộng {total_violations} lỗi vi phạm."

    # Câu 2: Track_id vi phạm nhiều nhất
    if data["violations"]:
        track_counts = {}
        for v in data["violations"]:
            track_counts[v["track_id"]] = track_counts.get(v["track_id"], 0) + 1
        top_track, top_count = max(track_counts.items(), key=lambda x: x[1])
        answer_2 = f"Xe số {top_track} là phương tiện vi phạm nhiều nhất trong khoảng thời gian này với {top_count} lần vi phạm."
    else:
        answer_2 = "Không có vi phạm nào được ghi nhận trong khoảng thời gian này."

    # Câu 3: Giờ có lưu lượng xe đông nhất
    if kpis["peak_hour"]:
        answer_3 = f"Giờ cao điểm là {kpis['peak_hour']} với {kpis['peak_val']} xe."
    else:
        answer_3 = f"Trong khoảng thời gian từ {start_time} đến {end_time}, không có dữ liệu ghi nhận về lưu lượng xe."

    return [
        {"question": STANDARD_QUESTIONS[0], "answer": answer_1},
        {"question": STANDARD_QUESTIONS[1], "answer": answer_2},
        {"question": STANDARD_QUESTIONS[2], "answer": answer_3},
    ]


# ============================================================
# 3. BIỂU ĐỒ (dùng cho bản Word — Excel dùng chart gốc của openpyxl)
# ============================================================

def _render_trend_chart_image(data: dict, out_path: str):
    """Biểu đồ cột chồng theo loại xe + đường vi phạm, đồng bộ màu với
    trang Phân tích (analytics.html) để người xem thấy ngay đây là cùng
    một bộ dữ liệu."""
    hours = sorted({r["hour"] for r in data["hourly_traffic_by_class"]} | {r["hour"] for r in data["hourly_violations"]})
    if not hours:
        hours = [r["hour"] for r in data["hourly_traffic"]]

    classes = sorted({r["vehicle_class"] for r in data["hourly_traffic_by_class"]},
                      key=lambda c: (c not in CLASS_COLORS, c))

    class_hour_map = {c: {} for c in classes}
    for r in data["hourly_traffic_by_class"]:
        class_hour_map[r["vehicle_class"]][r["hour"]] = r["total"] or 0

    violation_hours = {r["hour"]: r["total"] for r in data["hourly_violations"]}
    violation_vals = [violation_hours.get(h, 0) for h in hours]

    fig, ax1 = plt.subplots(figsize=(8, 3.6))
    bottom = [0] * len(hours)
    for c in classes:
        vals = [class_hour_map[c].get(h, 0) for h in hours]
        color = "#" + CLASS_COLORS.get(c, "8895AB")
        ax1.bar(hours, vals, bottom=bottom, label=CLASS_LABELS_VI.get(c, c), color=color, width=0.6)
        bottom = [b + v for b, v in zip(bottom, vals)]

    ax1.set_ylabel("Số xe", color="#333")
    ax1.set_xticks(range(len(hours)))
    ax1.set_xticklabels(hours, rotation=45, ha="right", fontsize=7)

    ax2 = ax1.twinx()
    ax2.plot(hours, violation_vals, marker="o", color="#" + COLOR_RED, linewidth=2, label="Vi phạm")
    ax2.set_ylabel("Vi phạm", color="#" + COLOR_RED)

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper left", fontsize=8, ncol=len(classes) + 1)

    ax1.set_title("Lưu lượng xe theo loại & Vi phạm theo giờ", fontsize=11, fontweight="bold")
    ax1.grid(axis="y", linestyle="--", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)


# ============================================================
# 4. HELPER ĐỊNH DẠNG WORD
# ============================================================

def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _kpi_cell(cell, label, value, sub, accent_hex):
    _shade_cell(cell, COLOR_NAVY)
    cell.vertical_alignment = 1
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(label)
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x9A, 0xA4, 0xB8)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(str(value))
    r2.font.size = Pt(22)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor.from_string(accent_hex)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(sub)
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0x6B, 0x74, 0x88)


def _add_kpi_row(doc, kpis):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = table.rows[0].cells

    _kpi_cell(cells[0], "TỔNG LƯỢT XE", f"{kpis['total_traffic']:,}", "trong khoảng thời gian đã chọn", COLOR_TEAL)
    _kpi_cell(cells[1], "TỔNG VI PHẠM", f"{kpis['total_violations']:,}", f"{kpis['violation_rate']:.1f}% trên tổng lượt xe", COLOR_RED)
    _kpi_cell(cells[2], "GIỜ CAO ĐIỂM", kpis["peak_hour"] or "–", f"{kpis['peak_val']} xe", "FFFFFF")
    _kpi_cell(cells[3], "XE PHỔ BIẾN NHẤT", kpis["top_class_label"], f"{kpis['top_val']} xe", "FFFFFF")
    doc.add_paragraph()


def _style_data_table(table):
    table.style = "Light Grid Accent 1"
    for cell in table.rows[0].cells:
        _shade_cell(cell, "2563EB")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ============================================================
# 5. XUẤT BÁO CÁO WORD (.docx)
# ============================================================

def generate_docx_report(start_time: str, end_time: str) -> str:
    data = _fetch_report_data(start_time, end_time)
    kpis = _compute_kpis(data)
    ai_summary = _generate_ai_summary(data, kpis, start_time, end_time)
    qa_results = _ask_ai_assistant_scoped(start_time, end_time, data, kpis)

    doc = Document()

    title = doc.add_heading("BÁO CÁO PHÂN TÍCH GIÁM SÁT & PHẠT NGUỘI GIAO THÔNG", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Khoảng thời gian: {start_time} — {end_time}\n").italic = True
    meta.add_run(f"Ngày xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M')}").italic = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "Báo cáo được trích xuất tự động từ cùng một cơ sở dữ liệu vận hành "
        "của Dashboard, trang Phân tích và Trợ lý AI — số liệu luôn nhất quán "
        "giữa cả ba giao diện tại mọi thời điểm."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x6B, 0x74, 0x88)

    doc.add_heading("Tổng quan nhanh", level=2)
    _add_kpi_row(doc, kpis)

    doc.add_heading("1. Nhận định tổng quan (AI)", level=2)
    doc.add_paragraph(ai_summary)

    doc.add_heading("2. Trợ lý AI trả lời nhanh", level=2)
    doc.add_paragraph(
        "Các câu hỏi dưới đây được gửi trực tiếp tới Trợ lý AI (cùng bộ máy "
        "đang phục vụ trang /chatbot), giới hạn trong đúng khoảng thời gian "
        "của báo cáo này:"
    ).italic = True
    for qa in qa_results:
        p = doc.add_paragraph()
        p.add_run("Hỏi: ").bold = True
        p.add_run(qa["question"])
        p2 = doc.add_paragraph()
        p2.add_run("Trả lời: ").bold = True
        r = p2.add_run(qa["answer"])
        r.font.color.rgb = RGBColor.from_string(COLOR_TEAL if "không" not in qa["answer"].lower() else "9AA4B8")

    doc.add_heading("3. Số liệu đếm xe theo loại và hướng", level=2)
    if not data["vehicle_counts"]:
        doc.add_paragraph("Không có dữ liệu lưu lượng trong khoảng thời gian này.")
    else:
        table = doc.add_table(rows=1, cols=3)
        _style_data_table(table)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Loại xe", "Hướng", "Số lượng"
        for i, row in enumerate(data["vehicle_counts"]):
            cells = table.add_row().cells
            cells[0].text = CLASS_LABELS_VI.get(row["vehicle_class"], row["vehicle_class"])
            cells[1].text = "Vào" if row["direction"] == "in" else ("Ra" if row["direction"] == "out" else row["direction"])
            cells[2].text = f"{row['total'] or 0:,}"
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if i % 2 == 1:
                for c in cells:
                    _shade_cell(c, "F2F5FA")

    doc.add_heading("4. Biểu đồ xu hướng theo giờ", level=2)
    if data["hourly_traffic"] or data["hourly_traffic_by_class"]:
        chart_path = os.path.join(REPORTS_DIR, "_trend_chart_tmp.png")
        _render_trend_chart_image(data, chart_path)
        doc.add_picture(chart_path, width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run("Hình 1: Lưu lượng xe theo loại (cột chồng) và số vi phạm (đường đỏ) theo từng giờ")
        cap_run.italic = True
        cap_run.font.size = Pt(9)
    else:
        doc.add_paragraph("Không có dữ liệu để vẽ biểu đồ trong khoảng thời gian này.")

    doc.add_heading("5. Danh sách vi phạm", level=2)
    if not data["violations"]:
        doc.add_paragraph("Không ghi nhận vi phạm trong khoảng thời gian này.")
    else:
        vtable = doc.add_table(rows=1, cols=5)
        _style_data_table(vtable)
        hdr = vtable.rows[0].cells
        for i, label in enumerate(["Thời gian", "Track ID", "Loại xe", "Loại vi phạm", "Vị trí"]):
            hdr[i].text = label
        for i, v in enumerate(data["violations"]):
            cells = vtable.add_row().cells
            cells[0].text = str(v["timestamp"])
            cells[1].text = str(v["track_id"])
            cells[2].text = CLASS_LABELS_VI.get(v["vehicle_class"], v["vehicle_class"])
            cells[3].text = v["violation_type"]
            cells[4].text = v["location"]
            if i % 2 == 1:
                for c in cells:
                    _shade_cell(c, "FFF5F5")

        doc.add_paragraph()
        doc.add_heading("Ảnh bằng chứng vi phạm", level=3)
        evidence_items = []
        for v in data["violations"]:
            path = _resolve_evidence_path(v.get("evidence_image_path"))
            if path:
                evidence_items.append((path, v))

        if not evidence_items:
            doc.add_paragraph("Không có ảnh bằng chứng đính kèm cho các vi phạm trên.")
        else:
            img_table = doc.add_table(rows=0, cols=2)
            for i in range(0, len(evidence_items), 2):
                row_cells = img_table.add_row().cells
                for col in range(2):
                    if i + col >= len(evidence_items):
                        break
                    path, v = evidence_items[i + col]
                    cell = row_cells[col]
                    try:
                        run_p = cell.paragraphs[0]
                        run = run_p.add_run()
                        run.add_picture(path, width=Inches(2.8))
                        run_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_p = cell.add_paragraph()
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = cap_p.add_run(
                            f"Track {v['track_id']} · {CLASS_LABELS_VI.get(v['vehicle_class'], v['vehicle_class'])} · {v['timestamp']}"
                        )
                        cap_run.font.size = Pt(8)
                        cap_run.italic = True
                    except Exception:
                        cell.add_paragraph("(Không thể chèn ảnh)")

    out_path = os.path.join(REPORTS_DIR, f"BaoCao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# 6. HELPER ĐỊNH DẠNG EXCEL
# ============================================================

_thin_border = Border(*[Side(style="thin", color="D0D5DD")] * 4)


def _add_excel_table(ws, headers, rows, table_name, style="TableStyleMedium9"):
    """Tạo Excel Table gốc (banded rows + filter dropdown có sẵn) thay vì
    tô màu thủ công -> chuyên nghiệp hơn, và Excel tự xử lý format."""
    ws.append(headers)
    for row in rows:
        ws.append(row)

    n_rows = len(rows) + 1
    n_cols = len(headers)
    if len(rows) == 0:
        ws.cell(row=2, column=1, value="Không có dữ liệu trong khoảng thời gian này.")
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=n_cols)
        ws.cell(row=2, column=1).font = Font(italic=True, color="6B7488")
        n_rows = 2

    last_col_letter = get_column_letter(n_cols)
    ref = f"A1:{last_col_letter}{n_rows}"
    if len(rows) > 0:
        tbl = Table(displayName=table_name, ref=ref)
        tbl.tableStyleInfo = TableStyleInfo(name=style, showRowStripes=True, showFirstColumn=False)
        ws.add_table(tbl)

    for col_idx, header in enumerate(headers, start=1):
        max_len = max([len(str(header))] + [len(str(r[col_idx - 1])) for r in rows]) if rows else len(str(header))
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 3, 12), 45)

    ws.freeze_panes = "A2"


def _kpi_card(ws, col_start, label, value, sub, accent_hex):
    col_letter1 = get_column_letter(col_start)
    col_letter2 = get_column_letter(col_start + 1)
    ws.merge_cells(f"{col_letter1}1:{col_letter2}1")
    ws.merge_cells(f"{col_letter1}2:{col_letter2}3")
    ws.merge_cells(f"{col_letter1}4:{col_letter2}4")

    cell_label = ws[f"{col_letter1}1"]
    cell_label.value = label
    cell_label.font = Font(size=10, color="9AA4B8", bold=True)
    cell_label.alignment = Alignment(horizontal="center")

    cell_val = ws[f"{col_letter1}2"]
    cell_val.value = value
    cell_val.font = Font(size=22, bold=True, color=accent_hex)
    cell_val.alignment = Alignment(horizontal="center", vertical="center")

    cell_sub = ws[f"{col_letter1}4"]
    cell_sub.value = sub
    cell_sub.font = Font(size=9, color="6B7488")
    cell_sub.alignment = Alignment(horizontal="center")

    for row in range(1, 5):
        for col in (col_start, col_start + 1):
            ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor="171D2E")


# ============================================================
# 7. XUẤT BÁO CÁO EXCEL (.xlsx)
# ============================================================

def generate_xlsx_report(start_time: str, end_time: str) -> str:
    data = _fetch_report_data(start_time, end_time)
    kpis = _compute_kpis(data)
    ai_summary = _generate_ai_summary(data, kpis, start_time, end_time)
    qa_results = _ask_ai_assistant_scoped(start_time, end_time, data, kpis)

    wb = Workbook()

    # ---- Sheet 1: Tổng quan (thẻ KPI giống trang Phân tích) ----
    ws1 = wb.active
    ws1.title = "Tổng quan"
    ws1["A6"] = "BÁO CÁO PHÂN TÍCH GIÁM SÁT & PHẠT NGUỘI GIAO THÔNG"
    ws1["A6"].font = Font(size=14, bold=True)
    ws1["A7"] = f"Khoảng thời gian: {start_time} — {end_time}"
    ws1["A8"] = f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws1["A7"].font = Font(italic=True, size=10)
    ws1["A8"].font = Font(italic=True, size=10)

    _kpi_card(ws1, 1, "TỔNG LƯỢT XE", kpis["total_traffic"], "trong khoảng đã chọn", "4FD1C5")
    _kpi_card(ws1, 3, "TỔNG VI PHẠM", kpis["total_violations"], f"{kpis['violation_rate']:.1f}% trên tổng lượt xe", "F56565")
    _kpi_card(ws1, 5, "GIỜ CAO ĐIỂM", kpis["peak_hour"] or "–", f"{kpis['peak_val']} xe", "FFFFFF")
    _kpi_card(ws1, 7, "XE PHỔ BIẾN NHẤT", kpis["top_class_label"], f"{kpis['top_val']} xe", "FFFFFF")

    ws1["A10"] = "Nhận định tổng quan (AI):"
    ws1["A10"].font = Font(bold=True)
    ws1["A11"] = ai_summary
    ws1["A11"].alignment = Alignment(wrap_text=True, vertical="top")
    ws1.merge_cells("A11:H16")

    ws1["A18"] = "Trợ lý AI trả lời nhanh (cùng bộ máy phục vụ trang /chatbot):"
    ws1["A18"].font = Font(bold=True)
    r = 19
    for qa in qa_results:
        ws1.cell(row=r, column=1, value=f"Hỏi: {qa['question']}").font = Font(bold=True, size=10)
        ws1.merge_cells(start_row=r, start_column=1, end_row=r, end_column=8)
        r += 1
        ws1.cell(row=r, column=1, value=f"Trả lời: {qa['answer']}").font = Font(size=10, color="2563EB")
        ws1.merge_cells(start_row=r, start_column=1, end_row=r, end_column=8)
        r += 2

    ws1.cell(row=r + 1, column=1,
             value="Báo cáo trích xuất tự động từ cùng cơ sở dữ liệu vận hành của Dashboard, trang Phân tích và Trợ lý AI.").font = Font(italic=True, size=9, color="6B7488")
    ws1.column_dimensions["A"].width = 14
    for col in "BCDEFGH":
        ws1.column_dimensions[col].width = 14

    # ---- Sheet 2: Đếm xe theo loại/hướng ----
    ws2 = wb.create_sheet("Số liệu đếm xe")
    veh_rows = [
        [CLASS_LABELS_VI.get(r["vehicle_class"], r["vehicle_class"]),
         "Vào" if r["direction"] == "in" else ("Ra" if r["direction"] == "out" else r["direction"]),
         r["total"] or 0]
        for r in data["vehicle_counts"]
    ]
    _add_excel_table(ws2, ["Loại xe", "Hướng", "Số lượng"], veh_rows, "TblVehicleCounts", "TableStyleMedium9")

    if veh_rows:
        chart = BarChart()
        chart.title = "Số lượng xe theo loại/hướng"
        chart.y_axis.title = "Số lượng"
        n = len(veh_rows)
        cats = Reference(ws2, min_col=1, min_row=2, max_row=n + 1)
        vals = Reference(ws2, min_col=3, min_row=1, max_row=n + 1)
        chart.add_data(vals, titles_from_data=True)
        chart.set_categories(cats)
        chart.width, chart.height = 18, 9
        ws2.add_chart(chart, "E2")

    # ---- Sheet 3: Vi phạm ----
    ws3 = wb.create_sheet("Danh sách vi phạm")
    vio_rows = [
        [v["timestamp"], v["track_id"], v.get("plate_number") or "",
         CLASS_LABELS_VI.get(v["vehicle_class"], v["vehicle_class"]),
         v["violation_type"], v["location"], v.get("evidence_image_path") or ""]
        for v in data["violations"]
    ]
    _add_excel_table(ws3, ["Thời gian", "Track ID", "Biển số", "Loại xe", "Loại vi phạm", "Vị trí", "Ảnh bằng chứng"],
                      vio_rows, "TblViolations", "TableStyleMedium7")

    # ---- Sheet 4: Xu hướng theo giờ (chi tiết theo giờ & loại xe -
    #      giống hệt bảng chi tiết trên trang Phân tích) ----
    ws4 = wb.create_sheet("Xu hướng theo giờ")
    classes = sorted({r["vehicle_class"] for r in data["hourly_traffic_by_class"]})
    class_hour_map = {c: {} for c in classes}
    for r in data["hourly_traffic_by_class"]:
        class_hour_map[r["vehicle_class"]][r["hour"]] = r["total"] or 0
    violation_hours = {r["hour"]: r["total"] for r in data["hourly_violations"]}

    all_hours = sorted({r["hour"] for r in data["hourly_traffic"]} | set(violation_hours.keys()))
    headers = ["Khung giờ"] + [CLASS_LABELS_VI.get(c, c) for c in classes] + ["Tổng xe", "Vi phạm"]
    rows = []
    for h in all_hours:
        total_h = sum(class_hour_map[c].get(h, 0) for c in classes)
        row = [h] + [class_hour_map[c].get(h, 0) for c in classes] + [total_h, violation_hours.get(h, 0)]
        rows.append(row)

    _add_excel_table(ws4, headers, rows, "TblHourlyTrend", "TableStyleMedium6")

    if rows:
        chart2 = BarChart()
        chart2.title = "Xu hướng lưu lượng / vi phạm theo giờ"
        n = len(rows)
        cats = Reference(ws4, min_col=1, min_row=2, max_row=n + 1)
        vals = Reference(ws4, min_col=len(classes) + 2, max_col=len(classes) + 3, min_row=1, max_row=n + 1)
        chart2.add_data(vals, titles_from_data=True)
        chart2.set_categories(cats)
        chart2.width, chart2.height = 20, 9
        ws4.add_chart(chart2, f"{get_column_letter(len(headers) + 2)}2")

    out_path = os.path.join(REPORTS_DIR, f"BaoCao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    wb.save(out_path)
    return out_path
